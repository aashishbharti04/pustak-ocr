"""Assemble corrected pages into DOCX, EPUB, or plain text."""

import html
import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from . import db

# Devanagari renders badly in Calibri; these ship with Windows and handle it properly.
BODY_FONT = "Nirmala UI"
BODY_SIZE = Pt(12)


def _paragraphs(text: str) -> list[str]:
    """Blank line = paragraph break, the usual convention for prose OCR."""
    blocks = re.split(r"\n\s*\n", text.strip())
    return [re.sub(r"\s*\n\s*", " ", b).strip() for b in blocks if b.strip()]


def _page_text(page) -> str:
    return (page["corrected_text"] or page["raw_ocr_text"] or "").strip()


def chapterize(book_id: int) -> list[dict]:
    """Group pages into chapters using the marks set during review.

    Pages before the first marked chapter become a leading untitled section, so a
    front-matter page never silently disappears from the export.
    """
    pages = db.list_pages(book_id)
    marks = {c["start_page"]: c["title"] for c in db.list_chapters(book_id)}

    sections: list[dict] = []
    current = {"title": None, "paragraphs": []}
    for page in pages:
        if page["page_number"] in marks:
            if current["paragraphs"] or current["title"]:
                sections.append(current)
            current = {"title": marks[page["page_number"]], "paragraphs": []}
        current["paragraphs"].extend(_paragraphs(_page_text(page)))

    if current["paragraphs"] or current["title"]:
        sections.append(current)
    return sections


def build_docx(book_id: int) -> io.BytesIO:
    book = db.get_book(book_id)
    if book is None:
        raise ValueError(f"no book {book_id}")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE

    title = doc.add_heading(book["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if book["author"]:
        byline = doc.add_paragraph(book["author"])
        byline.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for index, section in enumerate(chapterize(book_id)):
        if section["title"]:
            if index > 0:
                doc.add_page_break()
            doc.add_heading(section["title"], level=1)
        for block in section["paragraphs"]:
            doc.add_paragraph(block)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def build_epub(book_id: int) -> io.BytesIO:
    import tempfile
    from pathlib import Path

    from ebooklib import epub

    book = db.get_book(book_id)
    if book is None:
        raise ValueError(f"no book {book_id}")

    out = epub.EpubBook()
    out.set_identifier(f"pustak-ocr-{book_id}")
    out.set_title(book["title"])
    out.set_language("hi")
    if book["author"]:
        out.add_author(book["author"])

    css = epub.EpubItem(
        uid="style",
        file_name="style/main.css",
        media_type="text/css",
        content=(
            'body { font-family: "Nirmala UI", "Noto Sans Devanagari", serif; '
            "line-height: 1.8; }\n"
            "h1 { text-align: center; margin: 2em 0 1em; }\n"
            "p { text-indent: 1.5em; margin: 0 0 .3em; }\n"
        ),
    )
    out.add_item(css)

    chapters = []
    for index, section in enumerate(chapterize(book_id), start=1):
        title = section["title"] or (book["title"] if index == 1 else f"Section {index}")
        body = [f"<h1>{html.escape(title)}</h1>"]
        body += [f"<p>{html.escape(p)}</p>" for p in section["paragraphs"]]

        item = epub.EpubHtml(title=title, file_name=f"chap_{index:03d}.xhtml", lang="hi")
        item.content = "\n".join(body)
        item.add_item(css)
        out.add_item(item)
        chapters.append(item)

    out.toc = tuple(chapters)
    out.add_item(epub.EpubNcx())
    out.add_item(epub.EpubNav())
    out.spine = ["nav", *chapters]

    # ebooklib only writes to a path, so round-trip through a temp file.
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "book.epub"
        epub.write_epub(str(path), out)
        return io.BytesIO(path.read_bytes())


def build_text(book_id: int) -> str:
    parts = []
    for section in chapterize(book_id):
        if section["title"]:
            parts.append(section["title"])
        parts.extend(section["paragraphs"])
    return "\n\n".join(parts)
